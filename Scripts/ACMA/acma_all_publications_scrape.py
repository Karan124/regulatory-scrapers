#!/usr/bin/env python3
"""
ACMA Publications Scraper
Scrapes all publications from https://www.acma.gov.au/publications
with deduplication, PDF extraction, OCR for image-based PDFs, and anti-bot measures.
"""

import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from bs4 import BeautifulSoup
import json
import os
import logging
from datetime import datetime
import time
import random
from urllib.parse import urljoin, urlparse
import re
from pathlib import Path
from io import BytesIO

# PDF processing
import PyPDF2
import fitz  # PyMuPDF for better PDF handling
from PIL import Image

# Excel/CSV processing
import pandas as pd

# Test imports and set availability flags
try:
    import pytesseract
    OCR_AVAILABLE = True
    print("✅ pytesseract available for OCR processing")
except ImportError:
    OCR_AVAILABLE = False
    print("⚠️ pytesseract not available. Image-based PDF OCR will be skipped.")

try:
    from docx import Document
    DOCX_AVAILABLE = True
    print("✅ python-docx available for Word document processing")
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not available. Word document extraction will be skipped.")

# Configuration
BASE_URL = "https://www.acma.gov.au"
PUBLICATIONS_URL = f"{BASE_URL}/publications"
MAX_PAGES = 2  # Set to None for first run (scrape all), or number for daily runs
OUTPUT_FILE = "data/acma_publications.json"
LOG_FILE = "data/acma_scraper.log"
REQUEST_DELAY = (1, 3)  # Random delay between requests (seconds)

# Ensure data directory exists
os.makedirs("data", exist_ok=True)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(LOG_FILE),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

class ACMAScraper:
    def __init__(self):
        self.session = self._create_session()
        self.scraped_urls = set()
        self.existing_data = self._load_existing_data()
        self._load_scraped_urls()
        
    def _create_session(self):
        """Create a persistent session with realistic headers and retry strategy."""
        session = requests.Session()
        
        # Retry strategy
        retry_strategy = Retry(
            total=3,
            backoff_factor=1,
            status_forcelist=[403, 429, 500, 502, 503, 504],
        )
        adapter = HTTPAdapter(max_retries=retry_strategy)
        session.mount("http://", adapter)
        session.mount("https://", adapter)
        
        # Realistic headers to avoid bot detection
        session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.9',
            'Accept-Encoding': 'gzip, deflate, br',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1',
            'Sec-Fetch-Dest': 'document',
            'Sec-Fetch-Mode': 'navigate',
            'Sec-Fetch-Site': 'none',
            'Cache-Control': 'max-age=0'
        })
        
        return session
    
    def _load_existing_data(self):
        """Load existing publications data."""
        if os.path.exists(OUTPUT_FILE):
            try:
                with open(OUTPUT_FILE, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    logger.info(f"Loaded {len(data)} existing publications")
                    return data
            except Exception as e:
                logger.error(f"Error loading existing data: {e}")
                return []
        return []
    
    def _load_scraped_urls(self):
        """Load URLs of already scraped publications."""
        for item in self.existing_data:
            if 'url' in item:
                self.scraped_urls.add(item['url'])
    
    def _save_data(self, data):
        """Save data to JSON file."""
        try:
            with open(OUTPUT_FILE, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            logger.info(f"Saved {len(data)} publications to {OUTPUT_FILE}")
        except Exception as e:
            logger.error(f"Error saving data: {e}")
    
    def _random_delay(self):
        """Add random delay between requests."""
        delay = random.uniform(REQUEST_DELAY[0], REQUEST_DELAY[1])
        time.sleep(delay)
    
    def _initialize_session(self):
        """Initialize session by visiting homepage first."""
        try:
            logger.info("Initializing session by visiting homepage...")
            response = self.session.get(BASE_URL, timeout=30)
            response.raise_for_status()
            logger.info("Session initialized successfully")
            self._random_delay()
            return True
        except Exception as e:
            logger.error(f"Error initializing session: {e}")
            return False
    
    def _get_page(self, url):
        """Get a page with error handling."""
        try:
            self._random_delay()
            logger.info(f"Fetching: {url}")
            response = self.session.get(url, timeout=30)
            response.raise_for_status()
            return response
        except Exception as e:
            logger.error(f"Error fetching {url}: {e}")
            return None
    
    def _test_ocr_availability(self):
        """Test if OCR is available and working."""
        if not OCR_AVAILABLE:
            return False
            
        try:
            version = pytesseract.get_tesseract_version()
            logger.info(f"Tesseract OCR version {version} available")
            return True
        except Exception as e:
            logger.warning(f"OCR not available: {e}")
            return False
    
    def _extract_text_from_pdf(self, pdf_url):
        """Extract text from PDF file with OCR fallback for image-based PDFs."""
        try:
            response = self._get_page(pdf_url)
            if not response:
                return ""
            
            pdf_bytes = response.content
            text = ""
            
            # First try PyPDF2 for text-based PDFs
            try:
                pdf_reader = PyPDF2.PdfReader(BytesIO(pdf_bytes))
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += page_text + "\n"
                
                # Check if we got meaningful text
                clean_text = re.sub(r'\s+', ' ', text).strip()
                if len(clean_text) > 50:  # Threshold for meaningful text
                    logger.info(f"Extracted {len(clean_text)} characters from text-based PDF: {pdf_url}")
                    return clean_text
                else:
                    logger.info(f"Minimal text found from PyPDF2 ({len(clean_text)} chars), trying OCR for: {pdf_url}")
            except Exception as e:
                logger.warning(f"PyPDF2 extraction failed for {pdf_url}: {e}")
            
            # Try OCR for image-based PDFs using PyMuPDF
            if not self._test_ocr_availability():
                logger.warning(f"OCR not available, cannot process image-based PDF: {pdf_url}")
                return re.sub(r'\s+', ' ', text).strip()  # Return whatever we got from PyPDF2
            
            try:
                logger.info(f"Attempting OCR extraction for: {pdf_url}")
                pdf_document = fitz.open(stream=pdf_bytes)
                ocr_text = ""
                
                for page_num in range(len(pdf_document)):
                    page = pdf_document[page_num]
                    
                    # Try to get text directly first
                    page_text = page.get_text()
                    if len(page_text.strip()) > 50:
                        ocr_text += page_text + "\n"
                        logger.info(f"Extracted text directly from page {page_num + 1}")
                    else:
                        # Convert page to image and OCR
                        try:
                            # Render page as image with higher DPI for better OCR
                            pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
                            img_data = pix.tobytes("png")
                            img = Image.open(BytesIO(img_data))
                            
                            logger.info(f"Rendering page {page_num + 1} as {img.size[0]}x{img.size[1]} image for OCR")
                            
                            # Perform OCR with pytesseract - use default config first
                            try:
                                page_ocr_text = pytesseract.image_to_string(img)
                                if page_ocr_text.strip():
                                    ocr_text += page_ocr_text + "\n"
                                    logger.info(f"OCR extracted {len(page_ocr_text.strip())} characters from page {page_num + 1}")
                                else:
                                    logger.info(f"OCR found no text on page {page_num + 1}")
                            except Exception as default_ocr_error:
                                logger.warning(f"Default OCR failed for page {page_num + 1}: {str(default_ocr_error)}")
                                # Try with specific config
                                try:
                                    page_ocr_text = pytesseract.image_to_string(
                                        img, 
                                        lang='eng',
                                        config='--oem 3 --psm 6'
                                    )
                                    if page_ocr_text.strip():
                                        ocr_text += page_ocr_text + "\n"
                                        logger.info(f"Config OCR extracted {len(page_ocr_text.strip())} characters from page {page_num + 1}")
                                except Exception as config_ocr_error:
                                    logger.warning(f"Config OCR also failed for page {page_num + 1}: {str(config_ocr_error)}")
                                    continue
                            
                        except Exception as ocr_error:
                            logger.warning(f"OCR image processing failed for page {page_num + 1}: {str(ocr_error)}")
                            continue
                
                pdf_document.close()
                
                if ocr_text.strip():
                    clean_ocr_text = re.sub(r'\s+', ' ', ocr_text).strip()
                    logger.info(f"Total OCR extraction: {len(clean_ocr_text)} characters from {pdf_url}")
                    return clean_ocr_text
                else:
                    logger.warning(f"OCR produced no text for: {pdf_url}")
                    
            except Exception as e:
                logger.error(f"OCR processing failed for {pdf_url}: {e}")
            
            # Return whatever text we managed to extract
            return re.sub(r'\s+', ' ', text).strip()
            
        except Exception as e:
            logger.error(f"Error processing PDF {pdf_url}: {e}")
            return ""
    
    def _extract_text_from_word(self, word_url):
        """Extract text from MS Word documents (.docx)."""
        if not DOCX_AVAILABLE:
            logger.warning(f"python-docx not available, skipping Word document: {word_url}")
            return ""
            
        try:
            response = self._get_page(word_url)
            if not response:
                return ""
            
            if word_url.lower().endswith('.docx'):
                try:
                    doc = Document(BytesIO(response.content))
                    text_parts = []
                    
                    # Extract paragraphs
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text.strip())
                    
                    # Extract tables
                    for table in doc.tables:
                        for row in table.rows:
                            row_cells = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_cells.append(cell.text.strip())
                            if row_cells:
                                text_parts.append(" | ".join(row_cells))
                    
                    full_text = "\n".join(text_parts)
                    logger.info(f"Extracted {len(full_text)} characters from Word document: {word_url}")
                    return full_text
                    
                except Exception as e:
                    logger.error(f"Error extracting from .docx file {word_url}: {e}")
                    return ""
            
            elif word_url.lower().endswith('.doc'):
                logger.warning(f"Legacy .doc format not supported: {word_url}")
                return ""
                
        except Exception as e:
            logger.error(f"Error processing Word document {word_url}: {e}")
            return ""
    
    def _extract_text_from_excel_csv(self, file_url):
        """Extract text from Excel/CSV files."""
        try:
            response = self._get_page(file_url)
            if not response:
                return ""
            
            file_content = BytesIO(response.content)
            
            # Try Excel first
            try:
                df = pd.read_excel(file_content, engine='openpyxl')
                text = df.to_string(index=False)
                logger.info(f"Extracted Excel data ({len(text)} chars) from: {file_url}")
                return text
            except:
                # Try CSV
                try:
                    file_content.seek(0)
                    df = pd.read_csv(file_content)
                    text = df.to_string(index=False)
                    logger.info(f"Extracted CSV data ({len(text)} chars) from: {file_url}")
                    return text
                except Exception as csv_e:
                    logger.error(f"Failed to parse as Excel or CSV {file_url}: {csv_e}")
                    return ""
                
        except Exception as e:
            logger.error(f"Error extracting Excel/CSV from {file_url}: {e}")
            return ""
    
    def _extract_embedded_links(self, content_text):
        """Extract relevant embedded links from content."""
        if not content_text:
            return []
            
        links = []
        url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
        found_urls = re.findall(url_pattern, content_text)
        
        # Filter out social media and marketing links
        exclude_domains = ['facebook.com', 'twitter.com', 'linkedin.com', 'youtube.com', 'instagram.com']
        
        for url in found_urls:
            try:
                domain = urlparse(url).netloc.lower()
                if not any(excluded in domain for excluded in exclude_domains):
                    links.append(url)
            except:
                continue
        
        return list(set(links))  # Remove duplicates
    
    def _extract_publication_details(self, pub_element):
        """Extract details from a single publication element."""
        try:
            # Extract title and URL
            title_element = pub_element.find('h3', class_='card-title')
            if not title_element:
                return None
            
            link_element = title_element.find_parent('a') or title_element.find('a')
            if not link_element:
                return None
            
            title = title_element.get_text(strip=True)
            relative_url = link_element.get('href')
            publication_url = urljoin(BASE_URL, relative_url)
            
            # Skip if already scraped
            if publication_url in self.scraped_urls:
                logger.info(f"Skipping already scraped publication: {title}")
                return None
            
            # Extract date
            date_element = pub_element.find('time')
            published_date = date_element.get('datetime') if date_element else ""
            
            # Extract type/theme
            type_element = pub_element.find('p', class_='publication-type')
            theme = type_element.get_text(strip=True) if type_element else ""
            
            logger.info(f"Found new publication: {title}")
            return {
                'title': title,
                'url': publication_url,
                'published_date': published_date,
                'theme': theme
            }
            
        except Exception as e:
            logger.error(f"Error extracting publication details: {e}")
            return None
    
    def _scrape_publication_content(self, pub_details):
        """Scrape full content from a publication page."""
        try:
            response = self._get_page(pub_details['url'])
            if not response:
                return pub_details
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Extract image if available
            image_url = ""
            img_elements = soup.find_all('img')
            for img in img_elements:
                if img.get('src') and not img.get('src').startswith('data:'):
                    image_url = urljoin(BASE_URL, img.get('src'))
                    break
            
            # Find ALL file links - improved logic for multiple files
            pdf_texts = []
            csv_excel_texts = []
            word_texts = []
            embedded_links = []
            processed_files = set()
            
            # Method 1: Look in the publications-files-wrapper (primary method for ACMA)
            pub_files_wrapper = soup.find('div', class_='publications-files-wrapper')
            if pub_files_wrapper:
                logger.info("Found publications-files-wrapper, scanning for files...")
                file_items = pub_files_wrapper.find_all('div', class_='field__item')
                logger.info(f"Found {len(file_items)} file items in publications wrapper")
                
                for item in file_items:
                    file_link = item.find('a', href=True)
                    if file_link:
                        href = file_link.get('href', '').strip()
                        if href:
                            full_url = urljoin(BASE_URL, href)
                            
                            # Skip duplicates
                            if full_url in processed_files:
                                continue
                                
                            # Get file title for better logging
                            title_div = item.find('div', class_='file-title')
                            file_title = title_div.get_text(strip=True) if title_div else 'Unknown'
                            
                            # Determine file type
                            file_extension = ''
                            if '.' in href:
                                file_extension = href.lower().split('.')[-1].split('?')[0].split('#')[0]
                            
                            if file_extension == 'pdf':
                                logger.info(f"Processing PDF from wrapper: '{file_title}' - {full_url}")
                                pdf_content = self._extract_text_from_pdf(full_url)
                                if pdf_content:
                                    pdf_texts.append(f"=== {file_title} ===\n{pdf_content}")
                                    embedded_links.extend(self._extract_embedded_links(pdf_content))
                                processed_files.add(full_url)
                            
                            elif file_extension in ['xlsx', 'xls', 'csv']:
                                logger.info(f"Processing Excel/CSV from wrapper: '{file_title}' - {full_url}")
                                excel_content = self._extract_text_from_excel_csv(full_url)
                                if excel_content:
                                    csv_excel_texts.append(f"=== {file_title} ===\n{excel_content}")
                                processed_files.add(full_url)
                            
                            elif file_extension == 'docx':
                                if DOCX_AVAILABLE:
                                    logger.info(f"Processing Word document from wrapper: '{file_title}' - {full_url}")
                                    word_content = self._extract_text_from_word(full_url)
                                    if word_content:
                                        word_texts.append(f"=== {file_title} ===\n{word_content}")
                                        embedded_links.extend(self._extract_embedded_links(word_content))
                                    processed_files.add(full_url)
                                else:
                                    logger.warning(f"python-docx not available, skipping Word document: {file_title}")
                            
                            elif file_extension == 'doc':
                                logger.warning(f"Legacy .doc file found but not supported: '{file_title}' - {full_url}")
            
            # Method 2: Fallback - scan all links on the page (for any missed files)
            logger.info("Scanning all page links for additional files...")
            file_links = soup.find_all('a', href=True)
            
            additional_files_found = 0
            for link in file_links:
                href = link.get('href', '').strip()
                if not href:
                    continue
                    
                full_url = urljoin(BASE_URL, href)
                
                # Skip if already processed
                if full_url in processed_files:
                    continue
                
                # Only process file links (not page links)
                if '.' not in href:
                    continue
                    
                # Determine file type
                file_extension = href.lower().split('.')[-1].split('?')[0].split('#')[0]
                
                # Only process known file types
                if file_extension not in ['pdf', 'docx', 'doc', 'xlsx', 'xls', 'csv']:
                    continue
                
                additional_files_found += 1
                link_text = link.get_text(strip=True) or f"Additional {file_extension.upper()}"
                
                if file_extension == 'pdf':
                    logger.info(f"Processing additional PDF: '{link_text}' - {full_url}")
                    pdf_content = self._extract_text_from_pdf(full_url)
                    if pdf_content:
                        pdf_texts.append(f"=== {link_text} ===\n{pdf_content}")
                        embedded_links.extend(self._extract_embedded_links(pdf_content))
                    processed_files.add(full_url)
                
                elif file_extension in ['xlsx', 'xls', 'csv']:
                    logger.info(f"Processing additional Excel/CSV: '{link_text}' - {full_url}")
                    excel_content = self._extract_text_from_excel_csv(full_url)
                    if excel_content:
                        csv_excel_texts.append(f"=== {link_text} ===\n{excel_content}")
                    processed_files.add(full_url)
                
                elif file_extension == 'docx':
                    if DOCX_AVAILABLE:
                        logger.info(f"Processing additional Word document: '{link_text}' - {full_url}")
                        word_content = self._extract_text_from_word(full_url)
                        if word_content:
                            word_texts.append(f"=== {link_text} ===\n{word_content}")
                            embedded_links.extend(self._extract_embedded_links(word_content))
                        processed_files.add(full_url)
                    else:
                        logger.warning(f"python-docx not available, skipping additional Word document: {link_text}")
                
                elif file_extension == 'doc':
                    logger.warning(f"Legacy .doc file found but not supported: '{link_text}' - {full_url}")
            
            if additional_files_found > 0:
                logger.info(f"Found {additional_files_found} additional files via page scan")
            
            # Combine all content with better separation and titles
            all_pdf_text = "\n\n--- NEXT PDF ---\n\n".join(pdf_texts) if pdf_texts else ""
            all_csv_excel_text = "\n\n--- NEXT FILE ---\n\n".join(csv_excel_texts) if csv_excel_texts else ""
            all_word_text = "\n\n--- NEXT DOCUMENT ---\n\n".join(word_texts) if word_texts else ""
            
            # Log comprehensive summary
            total_files = len(pdf_texts) + len(csv_excel_texts) + len(word_texts)
            logger.info(f"Content extraction complete - Total files: {total_files} (PDFs: {len(pdf_texts)}, Excel/CSV: {len(csv_excel_texts)}, Word: {len(word_texts)})")
            
            if len(pdf_texts) > 1:
                logger.info(f"Successfully processed {len(pdf_texts)} PDF files")
            
            # Update publication details
            pub_details.update({
                'headline': pub_details['title'],
                'scraped_date': datetime.now().isoformat(),
                'image_url': image_url,
                'pdf_text': all_pdf_text,
                'csv_excel_text': all_csv_excel_text,
                'word_text': all_word_text,
                'embedded_links': list(set(embedded_links)),
                'files_processed': total_files  # Add file count for verification
            })
            
            return pub_details
            
        except Exception as e:
            logger.error(f"Error scraping publication content from {pub_details['url']}: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return pub_details
    
    def scrape_publications_page(self, page_num=0):
        """Scrape publications from a specific page."""
        url = f"{PUBLICATIONS_URL}?page={page_num}" if page_num > 0 else PUBLICATIONS_URL
        response = self._get_page(url)
        
        if not response:
            return [], False
        
        soup = BeautifulSoup(response.content, 'html.parser')
        publications = []
        
        # Find publication cards
        pub_cards = soup.find_all('article', class_='card-type-publication')
        logger.info(f"Found {len(pub_cards)} publication cards on page {page_num + 1}")
        
        for card in pub_cards:
            pub_details = self._extract_publication_details(card)
            if pub_details:
                # Scrape full content
                full_details = self._scrape_publication_content(pub_details)
                publications.append(full_details)
                
                # Add to scraped URLs to avoid duplicates
                self.scraped_urls.add(full_details['url'])
        
        # Check if there's a next page
        next_page_link = soup.find('li', class_='pager__item--next')
        has_next = next_page_link is not None
        
        logger.info(f"Processed {len(publications)} new publications from page {page_num + 1}")
        return publications, has_next
    
    def scrape_all_publications(self):
        """Main scraping function."""
        logger.info("=== Starting ACMA Publications Scraper ===")
        
        # Log availability of optional features
        if OCR_AVAILABLE:
            logger.info("✅ OCR available for image-based PDFs")
        else:
            logger.warning("⚠️ OCR not available - image-based PDFs will be skipped")
            
        if DOCX_AVAILABLE:
            logger.info("✅ Word document processing available")
        else:
            logger.warning("⚠️ Word document processing not available")
        
        # Initialize session
        if not self._initialize_session():
            logger.error("Failed to initialize session")
            return
        
        all_publications = list(self.existing_data)  # Start with existing data
        page_num = 0
        
        while True:
            # Check page limit for daily runs
            if MAX_PAGES is not None and page_num >= MAX_PAGES:
                logger.info(f"Reached maximum pages limit ({MAX_PAGES})")
                break
            
            logger.info(f"=== Scraping page {page_num + 1} ===")
            
            try:
                publications, has_next = self.scrape_publications_page(page_num)
                
                if not publications and page_num == 0:
                    logger.error("No publications found on first page. Site might be blocking requests.")
                    break
                
                # Add new publications
                new_count = 0
                for pub in publications:
                    # Check if already exists
                    if not any(existing['url'] == pub['url'] for existing in all_publications):
                        all_publications.append(pub)
                        new_count += 1
                
                logger.info(f"Added {new_count} new publications from page {page_num + 1}")
                
                # Save progress after each page
                self._save_data(all_publications)
                
                if not has_next:
                    logger.info("No more pages available")
                    break
                
                page_num += 1
                
            except Exception as e:
                logger.error(f"Error scraping page {page_num + 1}: {e}")
                break
        
        logger.info(f"=== Scraping completed. Total publications: {len(all_publications)} ===")
        self._save_data(all_publications)

def main():
    try:
        scraper = ACMAScraper()
        scraper.scrape_all_publications()
    except KeyboardInterrupt:
        logger.info("Scraping interrupted by user")
    except Exception as e:
        logger.error(f"Unexpected error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()